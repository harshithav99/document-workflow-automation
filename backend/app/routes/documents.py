import os
import shutil
import io
import PyPDF2
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
import hashlib
from sqlalchemy.orm import Session
from typing import List, Optional
from ..database import SessionLocal
from .. import models, schemas, auth
from datetime import datetime
from sqlalchemy import or_
import google.generativeai as genai
import docx

router = APIRouter(prefix="/documents", tags=["Documents"])

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

genai.configure(api_key="AIzaSyDGKag3Q6kkYPFHwK4y8eN0dnTD8BGSP-Y")

def log_audit(db, current_user, action: str, document_name: str, details: str = ""):
    u_name = getattr(current_user, "username", f"User {current_user.id}")
    doc_type = "pdf" if document_name.lower().endswith(".pdf") else "word"
    
    new_log = models.AuditLog(
        user_name=u_name,
        action=action,
        details=details,
        document_name=document_name,
        document_type=doc_type
    )
    db.add(new_log)
    db.commit()

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def generate_ai_summary(extracted_text: str) -> str:
    if not extracted_text or extracted_text == "No text extracted.":
        return "No text available to summarize."
    try:
        model = genai.GenerativeModel('gemini-2.5-flash')
        prompt = f"Please provide a concise, professional, 2-3 sentence summary of the following document:\n\n{extracted_text}"
        
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"AI Summarization failed: {e}")
        return "AI Summary could not be generated at this time."

@router.post("/", response_model=schemas.DocumentResponse)
async def upload_document(
    due_date_str: Optional[str] = Form(None),
    title: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    file_bytes = await file.read()
    document_hash = hashlib.sha256(file_bytes).hexdigest()
    
    safe_filename = file.filename.lower()
    if not safe_filename.endswith((".pdf", ".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are allowed.")

    extracted_text = ""
    print(f"\n--- STARTING EXTRACTION FOR: {file.filename} ---")
    
    if safe_filename.endswith('.pdf'):
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_text += page_text + "\n"
            print("PDF Extraction Successful!")
        except Exception as e:
            print(f"PDF ERROR: {repr(e)}") 
            extracted_text = "No text extracted."
            
    elif safe_filename.endswith('.docx'):
        try:
            word_doc = docx.Document(io.BytesIO(file_bytes))
            for para in word_doc.paragraphs:
                if para.text.strip():
                    extracted_text += para.text + "\n"
            for table in word_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            extracted_text += cell.text + " "
            print("DOCX Extraction Successful!")
        except Exception as e:
            print(f"DOCX ERROR: {repr(e)}") 
            extracted_text = "No text extracted."

    print(f"--- EXTRACTED {len(extracted_text)} CHARACTERS ---\n")

    final_summary = "No text extracted."
    if extracted_text and extracted_text != "No text extracted.":
        final_summary = generate_ai_summary(extracted_text[:30000])

    os.makedirs(UPLOAD_DIR, exist_ok=True)
    file_location = os.path.join(UPLOAD_DIR, file.filename)
    
    with open(file_location, "wb") as buffer:
        buffer.write(file_bytes)

    parsed_due_date = None
    if due_date_str:
        try:
            parsed_due_date = datetime.fromisoformat(due_date_str)
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid date format.")

    new_doc = models.Document(
        title=title,
        status="Draft",
        owner_id=current_user.id,
        owner_name=current_user.username,
        file_hash=document_hash,
        file_path=file_location, 
        summary=final_summary,
        due_date=parsed_due_date
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    log_audit(db, current_user, "Uploaded Document", new_doc.title)

    return new_doc
    
@router.get("/", response_model=List[schemas.DocumentResponse])
def get_documents(
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(auth.get_current_user)
):
    docs = db.query(models.Document).all()
    result = []
    
    for doc in docs:
        doc_data = {column.name: getattr(doc, column.name) for column in doc.__table__.columns}
        owner = db.query(models.User).filter(models.User.id == doc.owner_id).first()
        if owner:
            doc_data["owner_name"] = owner.username
        doc_data["assigned_reviewer_id"] = None
        doc_data["reviewer_name"] = None

        manual_id = getattr(doc, 'reviewer_id', None)

        if manual_id:
            doc_data["assigned_reviewer_id"] = manual_id
            reviewer = db.query(models.User).filter(models.User.id == manual_id).first()
            if reviewer:
                doc_data["reviewer_name"] = reviewer.username

        else:
            active_tasks = db.query(models.Task).filter(
                models.Task.document_id == doc.id,
                models.Task.status == "Pending"
            ).all()

            if len(active_tasks) > 1:
                doc_data["reviewer_name"] = "All Reviewers"
                # Unlock buttons if the current user is in the task list
                if any(task.reviewer_id == current_user.id for task in active_tasks):
                    doc_data["assigned_reviewer_id"] = current_user.id
                    
            elif len(active_tasks) == 1:
                doc_data["assigned_reviewer_id"] = active_tasks[0].reviewer_id
                reviewer = db.query(models.User).filter(models.User.id == active_tasks[0].reviewer_id).first()
                if reviewer:
                    doc_data["reviewer_name"] = reviewer.username

        result.append(doc_data)

    return result
@router.get("/my-tasks", response_model=List[schemas.DocumentResponse])
def get_my_tasks(
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(auth.get_current_user)
):
    user_role = str(current_user.role).strip().lower()

    my_pending_tasks = db.query(models.Task).filter(
        models.Task.reviewer_id == current_user.id,
        models.Task.status == "Pending" 
    ).all()
    
    assigned_doc_ids = [task.document_id for task in my_pending_tasks]
    
    if not assigned_doc_ids:
        assigned_doc_ids = [-1] 

    if "both" in user_role:
        docs = db.query(models.Document).filter(
            or_(
                models.Document.owner_id == current_user.id,
                models.Document.id.in_(assigned_doc_ids)
            )
        ).all()
        
    elif "author" in user_role:
        docs = db.query(models.Document).filter(models.Document.owner_id == current_user.id).all()
        
    elif "reviewer" in user_role:
        docs = db.query(models.Document).filter(models.Document.id.in_(assigned_doc_ids)).all()
        
    else:
        docs = []

    
    #for doc in docs:
        #doc_data = {column.name: getattr(doc, column.name) for column in doc.__table__.columns}
        #owner = db.query(models.User).filter(models.User.id == doc.owner_id).first()
        #doc_data["owner_name"] = owner.username if owner else "Unknown"
        #result.append(doc_data)
    result = []
    for doc in docs:
        doc_data = {column.name: getattr(doc, column.name) for column in doc.__table__.columns}
        
        owner = db.query(models.User).filter(models.User.id == doc.owner_id).first()
        doc_data["owner_name"] = (owner.full_name or owner.username) if owner else "Unknown"
        
        doc_data["assigned_date"] = getattr(doc, 'createdDate', None)
        
        user_task = db.query(models.Task).filter(
            models.Task.document_id == doc.id,
            models.Task.reviewer_id == current_user.id,
            models.Task.status == "Pending"
        ).first()
        
        if user_task and hasattr(user_task, 'createdDate'): 
            doc_data["assigned_date"] = user_task.createdDate

        pending_tasks = db.query(models.Task).filter(
            models.Task.document_id == doc.id,
            models.Task.status == "Pending"
        ).all()
        
        if len(pending_tasks) == 1:
            reviewer = db.query(models.User).filter(models.User.id == pending_tasks[0].reviewer_id).first()
            reviewer_name = (reviewer.full_name or reviewer.username) if reviewer else "Unknown"
        elif len(pending_tasks) > 1:
            reviewer_name = "All Reviewers"
        else:
            reviewer_name = "Unassigned"
            
        doc_data["reviewer_name"] = reviewer_name

        result.append(doc_data)
        
    return result

@router.get("/{doc_id}/download")
def download_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
    
    if not doc or not doc.file_path:
        raise HTTPException(status_code=404, detail="Document record not found")
        
    if not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="Physical file missing from server")

    return FileResponse(path=doc.file_path, filename=os.path.basename(doc.file_path))




@router.get("/{document_id}", response_model=schemas.DocumentResponse)
def get_document(
    document_id: int, 
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    doc_data = {column.name: getattr(doc, column.name) for column in doc.__table__.columns}
    owner = db.query(models.User).filter(models.User.id == doc.owner_id).first()
    if owner:
        doc_data["owner_name"] = owner.username
    doc_data["assigned_reviewer_id"] = None
    doc_data["reviewer_name"] = None

    manual_id = getattr(doc, 'reviewer_id', None)

    if manual_id:
        doc_data["assigned_reviewer_id"] = manual_id
        reviewer = db.query(models.User).filter(models.User.id == manual_id).first()
        if reviewer:
            doc_data["reviewer_name"] = reviewer.username

    else:
        active_tasks = db.query(models.Task).filter(
            models.Task.document_id == doc.id, 
            models.Task.status == "Pending"
        ).all()

        if len(active_tasks) > 1:
            doc_data["reviewer_name"] = "All Reviewers"
            if any(task.reviewer_id == current_user.id for task in active_tasks):
                doc_data["assigned_reviewer_id"] = current_user.id
                
        elif len(active_tasks) == 1:
            doc_data["assigned_reviewer_id"] = active_tasks[0].reviewer_id
            reviewer = db.query(models.User).filter(models.User.id == active_tasks[0].reviewer_id).first()
            if reviewer:
                doc_data["reviewer_name"] = reviewer.username

    return doc_data

@router.patch("/{document_id}/status", response_model=schemas.DocumentResponse)
async def update_document_status(
    document_id: int,
    status_update: schemas.StatusUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if status_update.status == "In Review":
        if current_user.role.lower() not in ["author", "both"]:
            raise HTTPException(status_code=403, detail="Only authors can send documents for review.")
        
        doc.status = "In Review"
        
        if status_update.reviewer_id:
            doc.assigned_reviewer_id = status_update.reviewer_id
            new_task = models.Task(
                document_id=doc.id,
                reviewer_id=status_update.reviewer_id,
                status="Pending"
            )
            db.add(new_task)
        else:
            doc.assigned_reviewer_id = None 
            
            all_reviewers = db.query(models.User).filter(
                or_(func.lower(User.role) == "reviewer", func.lower(User.role) == "both")
            ).all()
            
            existing = db.query(models.Task).filter(
                models.Task.document_id == doc.id, 
                models.Task.status == "Pending"
            ).first()
            
            if not existing:
                for reviewer in all_reviewers:
                    new_task = models.Task(
                        document_id=doc.id,
                        reviewer_id=reviewer.id,
                        status="Pending"
                    )
                    db.add(new_task)

    elif status_update.status in ["Approved", "Rejected"]:
        if current_user.role.lower() not in ["reviewer", "both"]:
            raise HTTPException(status_code=403, detail="Only reviewers can approve or reject.")
        
        active_task = db.query(models.Task).filter(
            models.Task.document_id == document_id,  
            models.Task.reviewer_id == current_user.id,
            models.Task.status == "Pending"
        ).first()

        if not active_task:
            raise HTTPException(
                status_code=403, 
                detail="Access Denied: You do not have a pending task for this document."
            )
            
        doc.status = "Draft" if status_update.status == "Rejected" else status_update.status
        active_task.status = "Completed"

        if status_update.status == "Rejected":
            doc.rejection_reason = getattr(status_update, 'reason', 'No reason provided')
            doc.due_date = None
        else:
            doc.rejection_reason = None        
        other_tasks = db.query(models.Task).filter(
            models.Task.document_id == document_id,
            models.Task.status == "Pending"
        ).all()
        for task in other_tasks:
            task.status = "Completed"

    new_log = models.AuditLog(
        user_name=current_user.username, 
        action=f"Document {status_update.status}", 
        details=getattr(status_update, 'reason', None) if status_update.status == "Rejected" else None,
        document_name=doc.title, 
    )
    db.add(new_log)

    db.commit()
    db.refresh(doc)
    
    return doc

@router.get("/audit/logs", response_model=list[schemas.AuditLogResponse])
async def get_audit_logs(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    logs = db.query(models.AuditLog).order_by(models.AuditLog.timestamp.desc()).all()
    return logs
@router.get("/debug/truth/{doc_id}")
def get_the_truth(doc_id: int, db: Session = Depends(get_db)):
    tasks = db.query(models.Task).filter(models.Task.document_id == doc_id).all()
    users = db.query(models.User).all()
    
    return {
        "tasks_for_document": len(tasks),
        "task_details": [{"task_id": t.id, "reviewer_id": t.reviewer_id, "status": t.status} for t in tasks],
        "every_user_in_database": [{"id": u.id, "username": u.username, "exact_role_text": u.role} for u in users]
    }

@router.put("/{document_id}/file", response_model=schemas.DocumentResponse)
async def upload_new_version(
    document_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    file_bytes = await file.read()
    document_hash = hashlib.sha256(file_bytes).hexdigest()

    extracted_snippet = "No text extracted."
    if file.filename.endswith('.pdf'):
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            if len(pdf_reader.pages) > 0:
                first_page = pdf_reader.pages[0]
                full_text = first_page.extract_text()
                extracted_snippet = full_text[:100].strip() 
        except Exception as e:
            print(f"Error reading PDF: {e}")

    file_location = f"{UPLOAD_DIR}/{file.filename}"
    with open(file_location, "wb") as buffer:
        buffer.write(file_bytes)

    doc.version += 1           
    doc.status = "Draft"       
    doc.file_hash = document_hash
    doc.file_path = file_location
    doc.summary = extracted_snippet

    db.commit()
    db.refresh(doc)

    log_audit(db, current_user, f"Uploaded Version {doc.version}.0", doc.title)
    return doc
@router.get("/notifications/user", response_model=list[schemas.AuditLogResponse])
async def get_user_notifications(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    user_role = current_user.role.lower()
    logs = []

    if user_role in ["author", "both"]:
        my_docs = db.query(models.Document).filter(models.Document.owner_id == current_user.id).all()
        my_doc_titles = [doc.title for doc in my_docs]
        
        if my_doc_titles:
            author_logs = db.query(models.AuditLog).filter(
                models.AuditLog.document_name.in_(my_doc_titles),
                models.AuditLog.action.in_(["Document Approved", "Document Rejected"]) 
            ).order_by(models.AuditLog.timestamp.desc()).limit(10).all()
            logs.extend(author_logs)

    if user_role in ["reviewer", "both"]:
        my_tasks = db.query(models.Task).filter(models.Task.reviewer_id == current_user.id).all()
        my_doc_ids = [task.document_id for task in my_tasks]
        
        if my_doc_ids:
            assigned_docs = db.query(models.Document).filter(models.Document.id.in_(my_doc_ids)).all()
            my_task_titles = [doc.title for doc in assigned_docs]
            
            if my_task_titles:
                reviewer_logs = db.query(models.AuditLog).filter(
                    models.AuditLog.document_name.in_(my_task_titles),
                    models.AuditLog.action == "Document In Review"
                ).order_by(models.AuditLog.timestamp.desc()).limit(10).all()
                logs.extend(reviewer_logs)

    unique_logs = {log.id: log for log in logs}.values()

    sorted_logs = sorted(unique_logs, key=lambda x: x.timestamp, reverse=True)[:10]
    
    return logs

@router.patch("/{doc_id}/deadline")
def update_document_deadline(
    doc_id: int,
    payload: schemas.DeadlineUpdate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    if doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Only the document owner can change the deadline.")

    try:
        new_deadline = datetime.fromisoformat(payload.due_date_str)
        doc.due_date = new_deadline
        db.commit()
        return {"message": "Deadline updated successfully!", "new_date": doc.due_date}
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format provided.")